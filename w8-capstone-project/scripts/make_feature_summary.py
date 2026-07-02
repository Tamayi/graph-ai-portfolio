"""Generate the feature summary (docs/SitRep_Feature_Summary.docx).

Embeds live-application screenshots captured by scripts/capture_screenshots.py.
Run after the screenshots exist:

    python scripts/make_feature_summary.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SHOTS = ROOT / "docs" / "assets" / "screenshots"
OUT = ROOT / "docs" / "SitRep_Feature_Summary_v2.docx"

TEAL = RGBColor(0x0B, 0x4F, 0x5C)
AMBER = RGBColor(0xD9, 0x7E, 0x00)
INK = RGBColor(0x1E, 0x2B, 0x2F)
MUTED = RGBColor(0x5E, 0x6E, 0x73)


def style_doc(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    for name, size in (("Heading 1", 16), ("Heading 2", 12.5)):
        st = doc.styles[name]
        st.font.name = "Georgia"
        st.font.size = Pt(size)
        st.font.color.rgb = TEAL
        st.font.bold = True


def p(doc, txt, color=INK, size=10.5, italic=False):
    par = doc.add_paragraph()
    r = par.add_run(txt)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.italic = italic
    return par


def screenshot(doc, name, caption):
    path = SHOTS / f"{name}.png"
    if not path.exists():
        p(doc, f"[ Place a screenshot here: save it as "
               f"docs/assets/screenshots/{name}.png, then rerun "
               "python scripts/make_feature_summary.py ]",
          color=AMBER, italic=True)
        return
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(str(path), width=Inches(6.9))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED
    r.font.italic = True


def feature(doc, title, body, shot, caption):
    doc.add_heading(title, level=1)
    p(doc, body)
    screenshot(doc, shot, caption)


def build():
    doc = Document()
    style_doc(doc)

    par = doc.add_paragraph()
    r = par.add_run("SITREP INTELLIGENCE — FEATURE SUMMARY")
    r.font.size = Pt(11); r.bold = True; r.font.color.rgb = AMBER
    par = doc.add_paragraph()
    r = par.add_run("The application, screen by screen")
    r.font.name = "Georgia"; r.font.size = Pt(24); r.bold = True
    r.font.color.rgb = TEAL
    p(doc, "All screenshots below are taken from the running application "
           "(FastAPI backend + web frontend on localhost). The interface is "
           "shown in English; the final section shows the same application in "
           "French to illustrate the multilingual design. All interface "
           "translations (FR / EN / PT) were generated with AI as part of the "
           "build.", color=MUTED)

    feature(doc, "1 · Overview dashboard",
            "The landing screen answers 'where does the outbreak stand today?' "
            "at a glance: cumulative confirmed cases, deaths, CFR%, affected "
            "health zones, patients in isolation and infected health workers — "
            "each with its 24-hour movement — plus the epidemic curve and "
            "per-province / per-zone breakdowns. Every number comes from the "
            "structured 'facts' datasets extracted from the reports, not from "
            "free text.",
            "overview_en", "Overview dashboard (English) — live data from the "
            "latest indexed SitRep.")

    feature(doc, "2 · Grounded chat assistant",
            "The assistant answers questions about the outbreak using RAG: it "
            "retrieves the most relevant report excerpts and must answer from "
            "them, citing SitRep number, date and page. For exact figures it "
            "calls the query_data tool (structured datasets), and for trends "
            "it calls plot_chart to render a chart in the conversation. The "
            "answer below was generated live by Claude.",
            "chat_en", "Chat assistant (English) — a grounded, cited answer "
            "with source references.")

    feature(doc, "3 · SitRep library & one-click ingestion",
            "The library lists every report in the corpus with its number, "
            "date and status. Adding a new report takes one action — paste an "
            "INSP page URL, a direct PDF link, or upload a file — and the "
            "pipeline does the rest: download, text extraction, number/date "
            "detection, AI extraction of structured facts, and semantic "
            "indexing. Progress is streamed step by step to the screen.",
            "library_en", "SitRep library (English) — the corpus, with "
            "URL/file ingestion at the top.")

    feature(doc, "4 · Single-report quality review",
            "One report is assessed the way a careful reviewer would read it: "
            "arithmetic (do provinces sum to the total?), internal consistency "
            "(do daily deaths reconcile with cumulative deaths?), timeliness, "
            "plausibility and geographic coverage. The result is an overall "
            "score, five dimension scores and a list of concrete issues, each "
            "tagged High / Medium / Low with its location in the document.",
            "dq_review_en", "Single review (English) — score, dimensions and "
            "severity-tagged issues.")

    feature(doc, "5 · Version comparison",
            "Two reports are validated against each other before their "
            "numbers are trusted: cumulative totals must never decrease, "
            "day-to-day jumps must be plausible, and missing values that "
            "break comparability are flagged. The screen shows each shared "
            "metric with its delta and an ok / info / warn / alert flag, "
            "plus section-level changes and a one-line verdict.",
            "dq_compare_en", "Version compare (English) — metric deltas with "
            "flags and findings.")

    feature(doc, "6 · Topics & gaps analysis",
            "The narrative 'défis et lacunes' of every report are extracted "
            "into themed items (financing, surveillance, vaccination, "
            "security, logistics, community engagement, human resources, "
            "laboratory) and synthesised across the whole series — showing "
            "which bottlenecks recur, which are worsening and which were "
            "resolved.",
            "topics_en", "Topics & gaps (English) — recurring themes across "
            "the report series.")

    doc.add_heading("7 · Multilingual by design", level=1)
    p(doc, "The source reports are French; the interface and the assistant "
           "work in French, English and Portuguese. Answers are grounded in "
           "the French source text and delivered in the user's language. The "
           "interface translations themselves were AI-generated during "
           "development — the same GenAI capability that powers the product "
           "also helped build it.")
    screenshot(doc, "overview_fr", "The same dashboard in French — one click "
               "on the language switcher.")
    screenshot(doc, "chat_fr", "The assistant answering in French, grounded "
               "in the same French source reports.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"saved {OUT}")


if __name__ == "__main__":
    build()
