"""Generate the capstone technical guide (docs/SitRep_Technical_Guide.docx).

Explains the five key processes for a mixed (non-developer) audience, with a
flow diagram per process rendered via matplotlib into docs/assets/. Run:

    python scripts/make_tech_guide.py
"""
from __future__ import annotations

from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets"
OUT = DOCS / "SitRep_Technical_Guide.docx"

# Palette shared with the overview deck
TEAL = "#0B4F5C"
TEAL_TINT = "#E7F0F1"
AMBER = "#D97E00"
INK = "#1E2B2F"
MUTED = "#5E6E73"
PAPER = "#FBFAF6"

TEAL_D = RGBColor(0x0B, 0x4F, 0x5C)
AMBER_D = RGBColor(0xD9, 0x7E, 0x00)
INK_D = RGBColor(0x1E, 0x2B, 0x2F)
MUTED_D = RGBColor(0x5E, 0x6E, 0x73)


# --- Diagram engine ----------------------------------------------------------
def flow_diagram(path: Path, rows: list[list[dict]], title: str | None = None,
                 figw: float = 11.0, figh: float = 2.4):
    """Render rows of boxes with arrows between consecutive boxes in a row.

    Each box: {t: title, s: subtitle, fill, edge, tc (title color)}.
    Multiple rows stack vertically (used for the two-store split and the
    two-phase RAG picture).
    """
    fig, ax = plt.subplots(figsize=(figw, figh), dpi=200)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 10 * len(rows))
    ax.axis("off")
    fig.patch.set_facecolor(PAPER)

    for ri, row in enumerate(rows):
        y = 10 * (len(rows) - 1 - ri) + 5  # row centre
        n = len(row)
        total_gap = 6 * (n - 1)
        bw = (96 - total_gap) / n
        x = 2
        for bi, b in enumerate(row):
            fill = b.get("fill", "white")
            edge = b.get("edge", "#E4DFD3")
            tc = b.get("tc", TEAL)
            sc = b.get("sc", MUTED)
            ax.add_patch(plt.Rectangle((x, y - 3.4), bw, 6.8, fc=fill, ec=edge,
                                       lw=1.2, joinstyle="round", zorder=2))
            ax.text(x + bw / 2, y + 1.1, b["t"], ha="center", va="center",
                    fontsize=10.5, fontweight="bold", color=tc, zorder=3,
                    family="DejaVu Sans")
            if b.get("s"):
                ax.text(x + bw / 2, y - 1.4, b["s"], ha="center", va="center",
                        fontsize=8.2, color=sc, zorder=3, family="DejaVu Sans",
                        wrap=True)
            if bi < n - 1:
                ax.annotate("", xy=(x + bw + 5.2, y), xytext=(x + bw + 0.8, y),
                            arrowprops=dict(arrowstyle="-|>", lw=2.2,
                                            color=AMBER, mutation_scale=18))
            x += bw + 6
    if title:
        ax.set_title(title, fontsize=11, color=INK, fontweight="bold",
                     family="DejaVu Sans", pad=10)
    fig.tight_layout()
    fig.savefig(path, facecolor=PAPER, bbox_inches="tight")
    plt.close(fig)


def make_diagrams() -> dict[str, Path]:
    ASSETS.mkdir(parents=True, exist_ok=True)
    d = {}

    d["library"] = ASSETS / "diagram_library.png"
    flow_diagram(d["library"], [[
        {"t": "1 · Add a report", "s": "INSP page URL, PDF\nlink or file upload"},
        {"t": "2 · Resolve & fetch", "s": "find the PDF link,\ndownload, verify %PDF"},
        {"t": "3 · Read the text", "s": "pypdf, page by page →\nMarkdown '## Page N'"},
        {"t": "4 · Register", "s": "archive PDF + Markdown,\nadd row to manifest.csv"},
        {"t": "5 · Make it usable", "s": "AI extraction + search\nindexing (next sections)",
         "fill": TEAL, "tc": "white", "sc": "#C8D9DC", "edge": TEAL},
    ]])

    d["rag"] = ASSETS / "diagram_rag.png"
    flow_diagram(d["rag"], [
        [
            {"t": "Markdown report", "s": "one file per SitRep"},
            {"t": "Chunking", "s": "split per page, then ~1400\nchars, 200 overlap"},
            {"t": "Voyage AI embeddings", "s": "voyage-4 multilingual,\none vector per chunk"},
            {"t": "Chroma index", "s": "vectors + SitRep N°,\ndate, page metadata",
             "fill": TEAL_TINT, "edge": TEAL_TINT},
        ],
        [
            {"t": "User question", "s": "any language"},
            {"t": "Embed the question", "s": "same Voyage model →\nsame vector space"},
            {"t": "Find nearest chunks", "s": "top 8 by cosine\nsimilarity in Chroma"},
            {"t": "Claude answers", "s": "grounded in excerpts,\ncited N° · date · page",
             "fill": TEAL, "tc": "white", "sc": "#C8D9DC", "edge": TEAL},
        ],
    ], figh=4.6)

    d["quality"] = ASSETS / "diagram_quality.png"
    flow_diagram(d["quality"], [[
        {"t": "One full report", "s": "complete Markdown,\nnot just excerpts"},
        {"t": "Claude reviews it", "s": "arithmetic, consistency,\nplausibility checks",
         "fill": TEAL, "tc": "white", "sc": "#C8D9DC", "edge": TEAL},
        {"t": "Structured verdict", "s": "forced JSON via a tool\ncall — never free prose"},
        {"t": "Score + issues", "s": "0-100, 5 dimensions,\nissues by severity",
         "fill": TEAL_TINT, "edge": TEAL_TINT},
    ]])

    d["compare"] = ASSETS / "diagram_compare.png"
    flow_diagram(d["compare"], [[
        {"t": "Report A + Report B", "s": "two full reports,\nearlier vs newer"},
        {"t": "Claude compares", "s": "shared metrics, deltas,\nreconciliation rules",
         "fill": TEAL, "tc": "white", "sc": "#C8D9DC", "edge": TEAL},
        {"t": "Flags per metric", "s": "ok · info · warn · alert\n(e.g. totals must not drop)"},
        {"t": "Verdict", "s": "issues by severity, section\nchanges, one-line banner",
         "fill": TEAL_TINT, "edge": TEAL_TINT},
    ]])

    d["topics"] = ASSETS / "diagram_topics.png"
    flow_diagram(d["topics"], [[
        {"t": "Every report", "s": "the 'défis & lacunes'\nnarrative sections"},
        {"t": "Extract per report", "s": "Claude → JSON items:\n{theme, detail}",
         "fill": TEAL, "tc": "white", "sc": "#C8D9DC", "edge": TEAL},
        {"t": "Aggregate", "s": "one list across the whole\nseries, cached as JSON"},
        {"t": "Synthesise", "s": "Claude → recurring themes,\ntrends, recommendations",
         "fill": TEAL, "tc": "white", "sc": "#C8D9DC", "edge": TEAL},
    ]])
    return d


# --- Document helpers ---------------------------------------------------------
def style_doc(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK_D
    for name, size, color in (("Heading 1", 17, TEAL_D),
                              ("Heading 2", 13, TEAL_D),
                              ("Heading 3", 11.5, INK_D)):
        st = doc.styles[name]
        st.font.name = "Georgia"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True


def p(doc, text="", bold_lead=None, size=10.5, color=INK_D, style=None):
    par = doc.add_paragraph(style=style)
    if bold_lead:
        r = par.add_run(bold_lead + "  ")
        r.bold = True
        r.font.color.rgb = TEAL_D
        r.font.size = Pt(size)
    r = par.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return par


def bullets(doc, items):
    for lead, body in items:
        par = doc.add_paragraph(style="List Bullet")
        r = par.add_run(lead + "  ")
        r.bold = True
        r.font.size = Pt(10.5)
        r2 = par.add_run(body)
        r2.font.size = Pt(10.5)


def ai_vs_code(doc, ai: str, code: str):
    """Callout every section repeats: what is GenAI vs plain programming."""
    t = doc.add_table(rows=2, cols=1)
    t.style = "Light Grid Accent 1"
    c0 = t.cell(0, 0).paragraphs[0]
    r = c0.add_run("Where the AI is:  ")
    r.bold = True; r.font.color.rgb = AMBER_D; r.font.size = Pt(10)
    r = c0.add_run(ai); r.font.size = Pt(10)
    c1 = t.cell(1, 0).paragraphs[0]
    r = c1.add_run("Where it is plain code:  ")
    r.bold = True; r.font.color.rgb = TEAL_D; r.font.size = Pt(10)
    r = c1.add_run(code); r.font.size = Pt(10)
    doc.add_paragraph()


def diagram(doc, path: Path, width=6.7):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(str(path), width=Inches(width))


# ==============================================================================
def build():
    imgs = make_diagrams()
    doc = Document()
    style_doc(doc)

    # Title block
    par = doc.add_paragraph()
    r = par.add_run("SITREP INTELLIGENCE — TECHNICAL GUIDE")
    r.font.size = Pt(11); r.bold = True; r.font.color.rgb = AMBER_D
    par = doc.add_paragraph()
    r = par.add_run("How the system works, process by process")
    r.font.name = "Georgia"; r.font.size = Pt(24); r.bold = True
    r.font.color.rgb = TEAL_D
    p(doc, "A grounded RAG application over the INSP situation reports of the "
           "17th Ebola (MVE/BVD) outbreak in DR Congo. This guide explains the "
           "five key processes at a high level — enough to understand what "
           "happens, in what order, and where generative AI is (and is not) "
           "involved. It is written for readers who are not developers.",
      color=MUTED_D)
    p(doc, "Models used: Anthropic Claude (Sonnet by default, Opus available "
           "for the hardest tasks) for reading, extraction, assessment and "
           "synthesis; Voyage AI (voyage-4, multilingual) for embeddings; "
           "ChromaDB as the local vector store.", color=MUTED_D)

    # --- 1. The SitRep library ------------------------------------------------
    doc.add_heading("1 · The SitRep library — adding a document", level=1)
    p(doc, "The library is the front door of the system. An analyst adds a "
           "report in one of three ways: an INSP web page URL, a direct PDF "
           "link, or a file upload. Everything else is automatic.")
    diagram(doc, imgs["library"])
    bullets(doc, [
        ("Resolve & fetch.", "If given a web page, the system scans it for "
         "the PDF link. The file is downloaded and verified to really be a "
         "PDF before anything else happens."),
        ("Read the text.", "pypdf extracts the text page by page. Each page "
         "becomes a '## Page N' section in a Markdown file — keeping page "
         "numbers is what makes page-level citations possible later."),
        ("Detect number & date.", "The SitRep number (N°) and report date are "
         "parsed from the URL, the filename or the document text; the analyst "
         "can override both."),
        ("Register.", "The PDF and Markdown files are archived and a row is "
         "added to manifest.csv — the library's simple, inspectable index."),
        ("Make it usable.", "The report is then structurally extracted with "
         "Claude and indexed for semantic search (sections 2 and 3). These "
         "steps are best-effort: if an API key is missing, the upload is "
         "still saved and can be reprocessed later."),
    ])
    ai_vs_code(doc,
        "only the last step — Claude turns the report text into a structured "
        "'facts' record (and the indexing uses AI embeddings).",
        "fetching, PDF-to-text conversion, number/date detection and the "
        "manifest are ordinary deterministic Python — cheap, fast, auditable.")

    # --- 2. RAG ----------------------------------------------------------------
    doc.add_heading("2 · RAG — how the chat finds and uses the right text", level=1)
    p(doc, "RAG (Retrieval-Augmented Generation) means the model does not "
           "answer from memory: for every question, the system first retrieves "
           "the most relevant passages from the actual reports, and the model "
           "must answer from those passages only. This is what makes answers "
           "grounded and citable.")
    doc.add_heading("Indexing (top row) — done once per report", level=2)
    bullets(doc, [
        ("Chunking.", "Each Markdown report is split on page boundaries, and "
         "any long page is further split into ~1,400-character pieces (about "
         "350-400 tokens) with a 200-character overlap, so a sentence cut at "
         "a boundary still appears whole in one of the chunks. Every chunk "
         "keeps its SitRep number, date and page."),
        ("Vectorisation with Voyage AI.", "Each chunk is sent to the voyage-4 "
         "embedding model, which returns a vector — a list of numbers that "
         "captures the meaning of the text. voyage-4 is multilingual, so "
         "French report text and an English question land in the same "
         "'meaning space'."),
        ("Storage.", "Vectors and their metadata are stored in a local "
         "ChromaDB collection using cosine similarity."),
    ])
    doc.add_heading("Question time (bottom row) — every chat turn", level=2)
    bullets(doc, [
        ("Retrieve.", "The question is embedded with the same model, and "
         "Chroma returns the 8 most similar chunks."),
        ("Answer.", "Claude receives the question plus those excerpts, with "
         "instructions to answer only from them and to cite SitRep N°, date "
         "and page for every figure. For exact numbers it can also call the "
         "query_data tool (section 6) instead of trusting narrative text."),
    ])
    diagram(doc, imgs["rag"], width=6.7)
    ai_vs_code(doc,
        "the embedding model (meaning-aware search) and Claude's grounded, "
        "cited answering, including its decision to call tools.",
        "chunking rules, the vector database, similarity search and citation "
        "bookkeeping are deterministic plumbing.")

    # --- 3. Single-report quality ------------------------------------------------
    doc.add_heading("3 · Single-SitRep data quality assessment", level=1)
    p(doc, "One report is reviewed the way a careful epidemiologist would "
           "review it — but in seconds. Claude receives the complete report "
           "(not retrieved snippets, so nothing is missed) together with a "
           "detailed checklist prompt.")
    diagram(doc, imgs["quality"])
    bullets(doc, [
        ("Arithmetic.", "Do province rows sum to the national total? Do "
         "health-zone counts add up? Is CFR% consistent with cases and "
         "deaths? Does the isolation balance work out?"),
        ("Consistency.", "Do daily deaths reconcile with the change in "
         "cumulative deaths? Does the narrative agree with the tables? Are "
         "missing values ('ND') breaking comparability?"),
        ("Plausibility.", "Operational red flags — sudden drops, backlog "
         "rebounds, contact-tracing saturation."),
        ("Output.", "The model must reply through a forced tool call that "
         "only accepts a fixed JSON shape: an overall score (0-100), five "
         "dimension scores (completeness, consistency, timeliness, "
         "plausibility, coverage), a list of issues each tagged High / "
         "Medium / Low with its location, and a one-sentence verdict."),
    ])
    ai_vs_code(doc,
        "the review itself — Claude does the reading, arithmetic reasoning "
        "and judgement, guided by a versioned prompt (quality_single_structured.md).",
        "the JSON schema enforcement, severity vocabulary and the UI that "
        "renders the score are fixed by the application, which is why the "
        "output is always displayable and comparable across reports.")

    # --- 4. Comparison -----------------------------------------------------------
    doc.add_heading("4 · Two-SitRep comparison", level=1)
    p(doc, "Outbreak numbers are cumulative: totals should never go down, and "
           "day-to-day jumps have to be plausible. The comparison module "
           "checks a new report against an earlier one before the numbers "
           "are trusted.")
    diagram(doc, imgs["compare"])
    bullets(doc, [
        ("Shared metrics.", "Claude lines up the indicators both reports "
         "share — confirmed cases, deaths, CFR%, affected health zones, "
         "isolation, contacts, laboratory — and computes the delta for each."),
        ("Flags.", "Every metric is flagged ok / info / warn / alert: a "
         "cumulative total that decreases is an alert; a suspiciously large "
         "jump is a warning; an expected change is ok."),
        ("Beyond numbers.", "It also reconciles deaths between the two dates, "
         "notes where 'ND' entries break comparability, and lists sections "
         "that were added, changed or dropped between the reports."),
        ("Output.", "Again a forced JSON shape: the metric table with flags, "
         "issues by severity, section changes, and a one-line banner verdict "
         "(e.g. how many errors were found)."),
    ])
    ai_vs_code(doc,
        "the cross-report reasoning — alignment of metrics that may be "
        "worded differently, judgement about what counts as implausible.",
        "the flag vocabulary, the JSON contract and the side-by-side UI.")

    # --- 5. Topics ---------------------------------------------------------------
    doc.add_heading("5 · Topic extraction — gaps & challenges over time", level=1)
    p(doc, "Every SitRep buries its most strategic content in prose: the "
           "'défis et lacunes' — gaps, constraints, unmet needs. Read one "
           "report and it is an anecdote; read forty and it is a pattern. "
           "This module makes the pattern visible.")
    diagram(doc, imgs["topics"])
    bullets(doc, [
        ("Extract (per report).", "Claude reads each report and returns JSON "
         "items {theme, detail}, using a fixed theme vocabulary: Financing, "
         "Surveillance, Vaccination, Security, Logistics, Community "
         "Engagement, Human Resources, Laboratory."),
        ("Aggregate.", "All items across the series are collected into one "
         "list (with SitRep number and date) and cached as JSON, so re-runs "
         "are instant."),
        ("Synthesise.", "A second Claude call reads the aggregate and writes "
         "the strategic narrative: which themes recur most (and in how many "
         "reports), which are emerging or worsening, which were resolved, "
         "and 3-5 prioritised recommendations."),
    ])
    ai_vs_code(doc,
        "both steps are generative — extraction of themes from prose, and the "
        "cross-series synthesis. This is the module that would be genuinely "
        "impossible without an LLM.",
        "the theme vocabulary, aggregation and caching.")

    # --- 6. Structured facts (supporting) ---------------------------------------
    doc.add_heading("6 · Under the hood: the structured 'facts' datasets", level=1)
    p(doc, "Several features above rely on one supporting process worth "
           "knowing about. When a report is added, Claude also extracts a "
           "structured record of its numbers — totals, per-province and "
           "per-health-zone tables, and response-pillar tables (laboratory, "
           "contacts, IPC, vaccination) — into a JSON file per report. Two "
           "design rules keep this honest:")
    bullets(doc, [
        ("Nullable by design.", "If a table is missing from the PDF, the "
         "field stays null. The model is explicitly forbidden from guessing."),
        ("Cached per report.", "Each report's extraction is saved "
         "individually; adding one new report re-extracts only that report, "
         "then re-aggregates the time-series datasets (KPI, province, zone, "
         "pillars) that power the dashboard and the chat's query_data tool."),
    ])

    # --- Appendix ----------------------------------------------------------------
    doc.add_heading("Appendix · Key settings at a glance", level=1)
    rows = [
        ("Chat / analysis model", "Anthropic Claude — claude-sonnet-4-6 "
         "default; claude-opus-4-8 for demanding tasks (configurable)"),
        ("Embedding model", "Voyage AI voyage-4 (multilingual)"),
        ("Vector store", "ChromaDB, local persistent collection 'insp_sitreps'"),
        ("Chunk size / overlap", "~1,400 characters (~350-400 tokens) / 200 characters"),
        ("Chunks retrieved per question", "top_k = 8"),
        ("Citations", "SitRep N° · report date · page, on every grounded figure"),
        ("Report archive", "data/sitreps/ — pdf/, md/, manifest.csv"),
        ("Structured datasets", "data/sitreps/datasets/ — per-report JSON + "
         "aggregated time-series"),
        ("Prompts", "prompts/*.md — externalised, versioned templates"),
        ("Token efficiency", "prompt caching on chat calls + per-turn history "
         "compaction (excerpts and tool payloads dropped once answered)"),
    ]
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        a, b = t.cell(i, 0), t.cell(i, 1)
        ra = a.paragraphs[0].add_run(k); ra.bold = True; ra.font.size = Pt(10)
        rb = b.paragraphs[0].add_run(v); rb.font.size = Pt(10)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"saved {OUT}")


if __name__ == "__main__":
    build()
